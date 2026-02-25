from fastapi import APIRouter, HTTPException, Depends, status
from fastapi.responses import StreamingResponse
from typing import Optional
import logging
import io

from app.models import PaperGenerationRequest, PaperGenerationResponse
from app.services.orchestration import get_orchestration_service
from app.services import get_cosmos_db_service
from app.utils import (
    ValidationError,
    AzureServiceError,
    AIAgentError,
    PaperGenerationError,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/papers", tags=["papers"])


@router.post("/generate", response_model=dict)
async def generate_paper(
    request: PaperGenerationRequest,
    orchestration_service=Depends(get_orchestration_service)
):
    """
    Generate a question paper
    
    Args:
        request: Paper generation request
        orchestration_service: Orchestration service
        
    Returns:
        Paper generation response with paper_id
    """
    try:
        # Validate request
        if not request.user_id:
            raise ValidationError("user_id is required")
        if not request.technology_topic or len(request.technology_topic.strip()) == 0:
            raise ValidationError("technology_topic is required")
        if request.num_questions < 1 or request.num_questions > 100:
            raise ValidationError("num_questions must be between 1 and 100")
        
        logger.info(
            f"Received paper generation request for user: {request.user_id}, "
            f"topic: {request.technology_topic}"
        )
        
        # Generate paper (returns paper_id)
        paper_id = await orchestration_service.generate_paper(request)
        
        return {
            "status": "queued",
            "paper_id": paper_id,
            "message": "Paper generation has been queued. Check status with the paper_id.",
        }
        
    except ValidationError as e:
        logger.warning(f"Validation error: {e.message}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"error_code": e.error_code, "message": e.message}
        )
    except AzureServiceError as e:
        logger.error(f"Azure service error: {e.message}")
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail={"error_code": e.error_code, "message": e.message}
        )
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )


@router.get("/status/{paper_id}")
async def get_paper_status(
    paper_id: str,
    cosmos_db_service=Depends(get_cosmos_db_service)
):
    """
    Get the status of a generated paper
    
    Args:
        paper_id: Paper ID
        cosmos_db_service: Cosmos DB service
        
    Returns:
        Paper status and metadata
    """
    try:
        # Retry logic: Cosmos DB query might be slow initially
        # Give it up to 10 seconds to find the record
        import asyncio
        max_retries = 20
        retry_delay = 0.5
        
        metadata = None
        for attempt in range(max_retries):
            metadata = await cosmos_db_service.get_by_id(paper_id)
            if metadata:
                break
            if attempt < max_retries - 1:
                await asyncio.sleep(retry_delay)
        
        if not metadata:
            # Paper not found after retries
            logger.warning(f"Paper not found after {max_retries} retries: {paper_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_NOT_FOUND", "message": f"Paper {paper_id} not found"}
            )
        
        return {
            "paper_id": paper_id,
            "status": metadata.get("status", "unknown"),
            "created_at": metadata.get("created_at"),
            "progress": metadata.get("progress", 0),
            "topic": metadata.get("topic"),
            "questions_count": metadata.get("questions_count"),
            "difficulty_distribution": metadata.get("difficulty_distribution"),
            "paper_url": metadata.get("paper_url"),
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting paper status: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )


@router.get("/user/{user_id}")
async def list_user_papers(
    user_id: str,
    cosmos_db_service=Depends(get_cosmos_db_service)
):
    """
    List all generated papers for a user
    
    Args:
        user_id: User ID
        cosmos_db_service: Cosmos DB service
        
    Returns:
        List of papers with metadata
    """
    try:
        query = "SELECT * FROM c WHERE c.user_id = @user_id AND c.topic != null ORDER BY c.created_at DESC"
        parameters = [{"name": "@user_id", "value": user_id}]
        papers = await cosmos_db_service.search_users(query, parameters)
        
        return {
            "user_id": user_id,
            "total": len(papers),
            "papers": [
                {
                    "paper_id": p.get("id"),
                    "topic": p.get("topic"),
                    "status": p.get("status"),
                    "created_at": p.get("created_at"),
                    "questions_count": p.get("questions_count"),
                    "difficulty_level": p.get("difficulty_level"),
                    "difficulty_distribution": p.get("difficulty_distribution"),
                    "duration_minutes": p.get("duration_minutes"),
                }
                for p in papers
            ],
        }
    except Exception as e:
        logger.error(f"Error listing papers for user {user_id}: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )


@router.get("/{paper_id}/content")
async def get_paper_content(
    paper_id: str,
    cosmos_db_service=Depends(get_cosmos_db_service)
):
    """
    Get the full paper content (questions, answers, etc.) from blob storage.
    Proxied through backend to avoid CORS issues.
    """
    from app.services import get_blob_storage_service
    try:
        metadata = await cosmos_db_service.get_by_id(paper_id)
        if not metadata:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_NOT_FOUND", "message": f"Paper {paper_id} not found"}
            )
        paper_url = metadata.get("paper_url")
        if not paper_url:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_CONTENT_NOT_FOUND", "message": "Paper content not available yet"}
            )
        # Extract blob name from the URL
        from urllib.parse import urlparse
        parsed = urlparse(paper_url)
        path_parts = parsed.path.lstrip("/").split("/", 1)
        blob_name = path_parts[1] if len(path_parts) > 1 else path_parts[0]
        blob_storage = await get_blob_storage_service()
        content = await blob_storage.download_paper(blob_name)
        return content
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting paper content: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )


@router.get("/{paper_id}/download")
async def download_paper_word(
    paper_id: str,
    cosmos_db_service=Depends(get_cosmos_db_service)
):
    """
    Download a generated paper as a Word (.docx) document.
    """
    from app.services import get_blob_storage_service
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import json as json_mod

    try:
        # Fetch metadata
        metadata = await cosmos_db_service.get_by_id(paper_id)
        if not metadata:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_NOT_FOUND", "message": f"Paper {paper_id} not found"}
            )
        paper_url = metadata.get("paper_url")
        if not paper_url:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_CONTENT_NOT_FOUND", "message": "Paper content not available yet"}
            )

        # Download blob content
        from urllib.parse import urlparse
        parsed_url = urlparse(paper_url)
        path_parts = parsed_url.path.lstrip("/").split("/", 1)
        blob_name = path_parts[1] if len(path_parts) > 1 else path_parts[0]
        blob_storage = await get_blob_storage_service()
        content = await blob_storage.download_paper(blob_name)

        # Parse questions
        questions = []
        if content.get("questions"):
            for q in content["questions"]:
                if q.get("question_text") or q.get("question"):
                    questions.append(q)
                elif q.get("raw"):
                    try:
                        cleaned = q["raw"]
                        if cleaned.startswith("```"):
                            cleaned = cleaned.split("\n", 1)[-1] if "\n" in cleaned else cleaned[3:]
                        if cleaned.endswith("```"):
                            cleaned = cleaned[:-3]
                        cleaned = cleaned.strip()
                        parsed_q = json_mod.loads(cleaned)
                        if isinstance(parsed_q, list):
                            questions.extend(parsed_q)
                        else:
                            questions.append(parsed_q)
                    except (json_mod.JSONDecodeError, Exception):
                        pass

        topic = content.get("topic", metadata.get("topic", "Interview Questions"))
        duration = content.get("duration_minutes", metadata.get("duration_minutes", ""))

        # Build Word document
        doc = Document()

        # Title
        title_heading = doc.add_heading(f"Interview Questions: {topic}", level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata line
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if duration:
            meta_run = meta_para.add_run(f"Duration: {duration} minutes  |  ")
            meta_run.font.size = Pt(11)
            meta_run.font.color.rgb = RGBColor(100, 100, 100)
        count_run = meta_para.add_run(f"Total Questions: {len(questions)}")
        count_run.font.size = Pt(11)
        count_run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph("")

        # Difficulty distribution
        diff_dist = content.get("difficulty_distribution", metadata.get("difficulty_distribution"))
        if diff_dist:
            dist_para = doc.add_paragraph()
            dist_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dist_run = dist_para.add_run("Difficulty Distribution: ")
            dist_run.bold = True
            dist_run.font.size = Pt(10)
            parts = [f"{level.capitalize()}: {count}" for level, count in diff_dist.items()]
            dist_detail = dist_para.add_run("  |  ".join(parts))
            dist_detail.font.size = Pt(10)
            dist_detail.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph("\u2500" * 60)

        # Difficulty badge colors
        difficulty_colors = {
            "easy": RGBColor(34, 139, 34),
            "medium": RGBColor(218, 165, 32),
            "hard": RGBColor(220, 20, 60),
        }

        # Questions
        for i, q in enumerate(questions, 1):
            q_text = q.get("question_text") or q.get("question", "")
            difficulty = (q.get("difficulty_level") or q.get("difficulty", "")).lower()

            # Question text
            q_para = doc.add_paragraph()
            q_num = q_para.add_run(f"Q{i}. ")
            q_num.bold = True
            q_num.font.size = Pt(12)
            q_body = q_para.add_run(q_text)
            q_body.font.size = Pt(12)

            # Difficulty badge
            if difficulty:
                badge = q_para.add_run(f"  [{difficulty.capitalize()}]")
                badge.font.size = Pt(9)
                badge.font.color.rgb = difficulty_colors.get(difficulty, RGBColor(100, 100, 100))
                badge.bold = True

            # Options
            options = q.get("options", [])
            if options:
                for j, opt in enumerate(options):
                    opt_para = doc.add_paragraph(style="List Bullet")
                    opt_run = opt_para.add_run(f"{chr(97 + j)}) {opt}")
                    opt_run.font.size = Pt(11)

            # Answer
            answer = q.get("correct_answer") or q.get("answer", "")
            if answer:
                ans_para = doc.add_paragraph()
                ans_label = ans_para.add_run("Answer: ")
                ans_label.bold = True
                ans_label.font.size = Pt(11)
                ans_label.font.color.rgb = RGBColor(0, 100, 0)
                ans_value = ans_para.add_run(str(answer))
                ans_value.font.size = Pt(11)

            # Explanation
            explanation = q.get("explanation", "")
            if explanation:
                exp_para = doc.add_paragraph()
                exp_label = exp_para.add_run("Explanation: ")
                exp_label.bold = True
                exp_label.font.size = Pt(10)
                exp_label.font.italic = True
                exp_value = exp_para.add_run(explanation)
                exp_value.font.size = Pt(10)
                exp_value.font.color.rgb = RGBColor(80, 80, 80)

            doc.add_paragraph("")

        # Save to buffer and return
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"interview_questions_{topic.replace(' ', '_').lower()}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating Word document: {e}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )


@router.get("/{paper_id}")
async def get_paper(
    paper_id: str,
    cosmos_db_service=Depends(get_cosmos_db_service)
):
    """
    Get a generated paper
    
    Args:
        paper_id: Paper ID
        cosmos_db_service: Cosmos DB service
        
    Returns:
        Paper content
    """
    from app.services import get_blob_storage_service
    try:
        metadata = await cosmos_db_service.get_by_id(paper_id)
        if not metadata:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={"error_code": "PAPER_NOT_FOUND", "message": f"Paper {paper_id} not found"}
            )
        # Generate SAS URL for the paper if paper_url exists
        paper_url = metadata.get("paper_url")
        sas_url = None
        if paper_url:
            # Extract blob name from the URL
            # Example: https://account.blob.core.windows.net/container/papers/paper_xxx/paper.json
            try:
                from urllib.parse import urlparse
                parsed = urlparse(paper_url)
                # Remove leading slash from path and remove container name
                path_parts = parsed.path.lstrip("/").split("/", 1)
                blob_name = path_parts[1] if len(path_parts) > 1 else path_parts[0]
                blob_storage = await get_blob_storage_service()
                sas_url = await blob_storage.get_sas_url(blob_name)
            except Exception as e:
                logger.warning(f"Failed to generate SAS URL for paper: {e}")
        return {
            "paper_id": paper_id,
            "topic": metadata.get("topic"),
            "status": metadata.get("status"),
            "created_at": metadata.get("created_at"),
            "questions_count": metadata.get("questions_count"),
            "difficulty_distribution": metadata.get("difficulty_distribution"),
            "paper_url": sas_url or paper_url,
            "user_id": metadata.get("user_id"),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting paper: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"error_code": "INTERNAL_ERROR", "message": str(e)}
        )
